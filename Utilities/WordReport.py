import os

from datetime import datetime

from docx import Document

from docx.shared import Inches


class WordReport:

    def __init__(self):

        project_path = os.path.dirname(os.path.dirname(__file__))

        self.report_path = os.path.join(
            project_path,
            "Reports",
            "Failed_TestCases.docx"
        )

        os.makedirs(
            os.path.dirname(self.report_path),
            exist_ok=True
        )

        self.document = Document()

        self.document.add_heading(
            "Automation Test Failure Report",
            level=1
        )

    def add_failure(
            self,
            testcase,
            expected,
            actual,
            reason,
            screenshot
    ):

        self.document.add_heading(
            testcase,
            level=2
        )

        self.document.add_paragraph(
            f"Execution Time : {datetime.now()}"
        )

        self.document.add_paragraph(
            f"Expected Result : {expected}"
        )

        self.document.add_paragraph(
            f"Actual Result : {actual}"
        )

        self.document.add_paragraph(
            f"Failure Reason : {reason}"
        )

        if os.path.exists(screenshot):

            self.document.add_picture(
                screenshot,
                width=Inches(5)
            )

        self.document.add_page_break()

    def save(self):

        self.document.save(self.report_path)